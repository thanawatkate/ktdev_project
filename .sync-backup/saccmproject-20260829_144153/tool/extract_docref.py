"""Extract content from document-ref folder to plain text for analysis."""
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from openpyxl import load_workbook
import docx
import pypdf

ROOT = Path(__file__).resolve().parents[1] / "docs" / "document-ref"
OUT = Path(__file__).resolve().parents[1] / "tool" / "_docref_extract.txt"


def dump_xlsx(p: Path, w):
    w.write(f"\n\n===== XLSX: {p.name} =====\n")
    try:
        wb = load_workbook(p, data_only=False, read_only=True)
    except Exception as e:
        w.write(f"[ERROR opening: {e}]\n")
        return
    for sn in wb.sheetnames:
        ws = wb[sn]
        w.write(f"\n--- Sheet: {sn} (max_row={ws.max_row}, max_col={ws.max_column}) ---\n")
        rows_written = 0
        for row in ws.iter_rows(values_only=True):
            if not any(c is not None and str(c).strip() != "" for c in row):
                continue
            line = " | ".join("" if c is None else str(c) for c in row)
            w.write(line + "\n")
            rows_written += 1
            if rows_written >= 80:
                w.write("... (truncated rows) ...\n")
                break
    wb.close()


def dump_docx(p: Path, w):
    w.write(f"\n\n===== DOCX: {p.name} =====\n")
    try:
        d = docx.Document(p)
    except Exception as e:
        w.write(f"[ERROR opening: {e}]\n")
        return
    for para in d.paragraphs:
        t = para.text.strip()
        if t:
            w.write(t + "\n")
    for ti, table in enumerate(d.tables):
        w.write(f"\n[Table #{ti+1}]\n")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            w.write(" | ".join(cells) + "\n")


def dump_pdf(p: Path, w, max_pages=80):
    w.write(f"\n\n===== PDF: {p.name} =====\n")
    try:
        reader = pypdf.PdfReader(str(p))
    except Exception as e:
        w.write(f"[ERROR opening: {e}]\n")
        return
    n = len(reader.pages)
    w.write(f"[total pages: {n}]\n")
    pages_to_read = min(n, max_pages)
    for i in range(pages_to_read):
        try:
            text = reader.pages[i].extract_text() or ""
        except Exception as e:
            text = f"[page extract error: {e}]"
        text = text.strip()
        if text:
            w.write(f"\n--- page {i+1} ---\n{text}\n")
    if n > pages_to_read:
        w.write(f"\n[... {n - pages_to_read} more pages not shown ...]\n")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    files = sorted(ROOT.rglob("*"))
    with OUT.open("w", encoding="utf-8") as w:
        w.write(f"# document-ref extraction\nroot: {ROOT}\n")
        for p in files:
            if not p.is_file():
                continue
            suf = p.suffix.lower()
            try:
                if suf == ".xlsx":
                    dump_xlsx(p, w)
                elif suf == ".docx":
                    dump_docx(p, w)
                elif suf == ".pdf":
                    dump_pdf(p, w)
            except Exception as e:
                w.write(f"\n[ERROR processing {p.name}: {e}]\n")
    print(f"Wrote: {OUT}")
    print(f"Size: {OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
